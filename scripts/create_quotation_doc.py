from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "documents"
LOGO = ROOT / "documents" / "frosty_lwt_logo_doc.jpg"
OUT = OUT_DIR / "Light Winter Technologies - Quotation - Anephen Investments - Address Fields.docx"


TEAL = "2A5F91"
DARK = "1D2733"
MUTED = "5B6776"
SOFT = "EEF5FA"
LINE = "C6D7E6"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=LINE, size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, size=10, color=DARK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_label_value(container, label, value):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor.from_string(DARK)
    r2 = p.add_run(value)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(MUTED)


def add_printable_watermark(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    watermark_xml = """
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xmlns:v="urn:schemas-microsoft-com:vml"
            xmlns:o="urn:schemas-microsoft-com:office:office">
      <v:shapetype id="_x0000_t136"
                   coordsize="21600,21600"
                   o:spt="136"
                   adj="10800"
                   path="m@7,l@8,m@5,21600l@6,21600e">
        <v:formulas>
          <v:f eqn="sum #0 0 10800"/>
          <v:f eqn="prod #0 2 1"/>
          <v:f eqn="sum 21600 0 @1"/>
          <v:f eqn="sum 0 0 @2"/>
          <v:f eqn="sum 21600 0 @3"/>
          <v:f eqn="if @0 @3 0"/>
          <v:f eqn="if @0 21600 @1"/>
          <v:f eqn="if @0 0 @2"/>
          <v:f eqn="if @0 @4 21600"/>
          <v:f eqn="mid @5 @6"/>
          <v:f eqn="mid @8 @5"/>
          <v:f eqn="mid @7 @8"/>
          <v:f eqn="mid @6 @7"/>
          <v:f eqn="sum @6 0 @5"/>
        </v:formulas>
        <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
        <v:textpath on="t" fitshape="t"/>
        <v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>
        <o:lock v:ext="edit" text="t" shapetype="t"/>
      </v:shapetype>
      <v:shape id="LightWinterWatermark"
               o:spid="_x0000_s1025"
               type="#_x0000_t136"
               fillcolor="#DDEAF5"
               stroked="f"
               o:allowincell="f"
               style="position:absolute;margin-left:-85pt;margin-top:210pt;width:620pt;height:100pt;rotation:315;z-index:-251654144;mso-position-horizontal:center;mso-position-horizontal-relative:page;mso-position-vertical:center;mso-position-vertical-relative:page">
        <v:fill opacity="0.38"/>
        <v:textpath style="font-family:Arial;font-size:30pt;font-weight:bold" string="LIGHT WINTER TECHNOLOGIES"/>
      </v:shape>
    </w:pict>
    """
    run._r.append(parse_xml(watermark_xml))


def main():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.45)
    section.right_margin = Cm(1.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after = Pt(2)
    if LOGO.exists():
        logo_para.add_run().add_picture(str(LOGO), width=Inches(0.82))

    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    header.columns[0].width = Cm(10.55)
    header.columns[1].width = Cm(6.2)
    left = header.cell(0, 0)
    right = header.cell(0, 1)
    for cell in (left, right):
        set_cell_border(cell, color="FFFFFF", size="0")

    title = left.paragraphs[0]
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("LIGHT WINTER TECHNOLOGIES")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(DARK)
    sub = left.add_paragraph()
    sub.paragraph_format.space_after = Pt(0)
    rs = sub.add_run("RetailOS Commercial Quotation")
    rs.font.size = Pt(9.5)
    rs.font.color.rgb = RGBColor.from_string(TEAL)

    q = right.paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rq = q.add_run("QUOTATION")
    rq.bold = True
    rq.font.size = Pt(24)
    rq.font.color.rgb = RGBColor.from_string(TEAL)
    add_label_value(right, "Quote No", "LWT-________")
    add_label_value(right, "Date", date.today().strftime("%d %B %Y"))
    add_label_value(right, "Valid Until", "________________")

    doc.add_paragraph()
    info = doc.add_table(rows=1, cols=2)
    info.autofit = False
    info.columns[0].width = Cm(8.3)
    info.columns[1].width = Cm(8.3)
    for cell in info.row_cells(0):
        shade_cell(cell, SOFT)
        set_cell_border(cell)
    set_cell_text(
        info.cell(0, 0),
        "Prepared For\n"
        "Anephen Investments\n"
        "Address: __________________________\n"
        "         __________________________\n"
        "Contact: __________________________\n"
        "Phone/Email: ______________________",
        bold=False,
        size=9.5,
    )
    set_cell_text(
        info.cell(0, 1),
        "Prepared By\n"
        "Light Winter Technologies\n"
        "Address: __________________________\n"
        "         __________________________\n"
        "Contact: __________________________\n"
        "Phone/Email: ______________________",
        bold=False,
        size=9.5,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    rt = p.add_run("Items")
    rt.bold = True
    rt.font.size = Pt(13)
    rt.font.color.rgb = RGBColor.from_string(DARK)

    table = doc.add_table(rows=6, cols=5)
    table.autofit = False
    widths = [Cm(1.1), Cm(7.5), Cm(2.0), Cm(2.8), Cm(3.0)]
    headings = ["#", "Description", "Qty", "Unit Price", "Total"]
    for i, (cell, heading) in enumerate(zip(table.rows[0].cells, headings)):
        cell.width = widths[i]
        shade_cell(cell, TEAL)
        set_cell_border(cell, color=TEAL)
        set_cell_text(cell, heading, bold=True, size=9.5, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    for row_idx in range(1, 6):
        cells = table.rows[row_idx].cells
        values = [str(row_idx), "", "", "", ""]
        for i, (cell, value) in enumerate(zip(cells, values)):
            cell.width = widths[i]
            set_cell_border(cell)
            set_cell_text(cell, value, size=10, align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT)

    totals = doc.add_table(rows=4, cols=2)
    totals.autofit = False
    totals.columns[0].width = Cm(12.8)
    totals.columns[1].width = Cm(3.8)
    rows = [("Subtotal", ""), ("Discount", ""), ("Tax / VAT", ""), ("Grand Total", "")]
    for r_idx, (label, value) in enumerate(rows):
        label_cell, value_cell = totals.rows[r_idx].cells
        for cell in (label_cell, value_cell):
            set_cell_border(cell)
            if r_idx == 3:
                shade_cell(cell, SOFT)
        set_cell_text(label_cell, label, bold=r_idx == 3, size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(value_cell, value, bold=r_idx == 3, size=10.5, align=WD_ALIGN_PARAGRAPH.RIGHT)

    notes_title = doc.add_paragraph()
    notes_title.paragraph_format.space_before = Pt(12)
    rn = notes_title.add_run("Notes")
    rn.bold = True
    rn.font.size = Pt(12)
    rn.font.color.rgb = RGBColor.from_string(DARK)
    notes = doc.add_table(rows=1, cols=1)
    set_cell_border(notes.cell(0, 0))
    shade_cell(notes.cell(0, 0), "F8FBFA")
    set_cell_text(notes.cell(0, 0), "______________________________________________________________________________\n______________________________________________________________________________", size=10, color=MUTED)

    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = False
    sig.columns[0].width = Cm(8.2)
    sig.columns[1].width = Cm(8.2)
    set_cell_text(sig.cell(0, 0), "Accepted By: __________________________\nDate: _________________________________", size=10)
    set_cell_text(sig.cell(0, 1), "For Light Winter Technologies\nSignature: ____________________________", size=10)
    for cell in sig.row_cells(0):
        set_cell_border(cell, color="FFFFFF", size="0")

    visible_watermark = doc.add_paragraph()
    visible_watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
    visible_watermark.paragraph_format.space_before = Pt(4)
    visible_watermark.paragraph_format.space_after = Pt(0)
    rw = visible_watermark.add_run("LIGHT WINTER TECHNOLOGIES")
    rw.bold = True
    rw.font.size = Pt(19)
    rw.font.color.rgb = RGBColor.from_string("D7E7F4")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run("Light Winter Technologies | Premium RetailOS Platform")
    rf.font.size = Pt(8.5)
    rf.font.color.rgb = RGBColor.from_string(MUTED)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
